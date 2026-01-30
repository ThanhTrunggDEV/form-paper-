"""
Style Applier for Springer LNCS Format
Applies Springer LNCS styling to documents
"""

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import Dict, List, Optional
import re
import os


class SpringerStyles:
    """Springer LNCS Style definitions"""
    
    # Font settings
    TITLE_FONT = 'Times New Roman'
    TITLE_SIZE = Pt(14)
    
    AUTHOR_FONT = 'Times New Roman'
    AUTHOR_SIZE = Pt(11)
    
    ABSTRACT_FONT = 'Times New Roman'
    ABSTRACT_SIZE = Pt(10)
    
    BODY_FONT = 'Times New Roman'
    BODY_SIZE = Pt(11)
    
    HEADING1_SIZE = Pt(12)
    HEADING2_SIZE = Pt(11)
    HEADING3_SIZE = Pt(10)
    
    CAPTION_FONT = 'Times New Roman'
    CAPTION_SIZE = Pt(10)
    
    # Page settings
    PAGE_WIDTH = Cm(16)  # A4 with margins
    PAGE_HEIGHT = Cm(24.7)
    MARGIN_TOP = Cm(2.5)
    MARGIN_BOTTOM = Cm(2.5)
    MARGIN_LEFT = Cm(2.5)
    MARGIN_RIGHT = Cm(2.5)
    
    # Spacing
    LINE_SPACING = 1.0  # Single spacing
    PARAGRAPH_AFTER = Pt(6)


class StyleApplier:
    """Apply Springer LNCS styles to documents"""
    
    def __init__(self):
        self.document = None
        self.styles = SpringerStyles()
        self.figure_counter = 0
        self.table_counter = 0
        self.changes_log = []
    
    def create_new_document(self) -> Document:
        """Create a new document with Springer styles"""
        self.document = Document()
        self._setup_page_layout()
        self._create_styles()
        return self.document
    
    def load_document(self, file_path: str) -> bool:
        """Load an existing document"""
        try:
            self.document = Document(file_path)
            return True
        except Exception as e:
            print(f"Error loading document: {e}")
            return False
    
    def _setup_page_layout(self):
        """Set up page margins and size"""
        section = self.document.sections[0]
        section.page_width = Cm(21)  # A4 width
        section.page_height = Cm(29.7)  # A4 height
        section.top_margin = self.styles.MARGIN_TOP
        section.bottom_margin = self.styles.MARGIN_BOTTOM
        section.left_margin = self.styles.MARGIN_LEFT
        section.right_margin = self.styles.MARGIN_RIGHT
        
        self.changes_log.append("Page margins set to 2.5cm all sides")
    
    def _create_styles(self):
        """Create custom styles for Springer LNCS"""
        styles = self.document.styles
        
        # Title style
        if 'LNCS Title' not in [s.name for s in styles]:
            title_style = styles.add_style('LNCS Title', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.name = self.styles.TITLE_FONT
            title_style.font.size = self.styles.TITLE_SIZE
            title_style.font.bold = True
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(12)
        
        # Author style
        if 'LNCS Author' not in [s.name for s in styles]:
            author_style = styles.add_style('LNCS Author', WD_STYLE_TYPE.PARAGRAPH)
            author_style.font.name = self.styles.AUTHOR_FONT
            author_style.font.size = self.styles.AUTHOR_SIZE
            author_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            author_style.paragraph_format.space_after = Pt(6)
        
        # Abstract style
        if 'LNCS Abstract' not in [s.name for s in styles]:
            abstract_style = styles.add_style('LNCS Abstract', WD_STYLE_TYPE.PARAGRAPH)
            abstract_style.font.name = self.styles.ABSTRACT_FONT
            abstract_style.font.size = self.styles.ABSTRACT_SIZE
            abstract_style.font.italic = True
            abstract_style.paragraph_format.space_after = Pt(12)
        
        # Caption style
        if 'LNCS Caption' not in [s.name for s in styles]:
            caption_style = styles.add_style('LNCS Caption', WD_STYLE_TYPE.PARAGRAPH)
            caption_style.font.name = self.styles.CAPTION_FONT
            caption_style.font.size = self.styles.CAPTION_SIZE
            caption_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_style.paragraph_format.space_before = Pt(6)
            caption_style.paragraph_format.space_after = Pt(12)
    
    def apply_springer_style(self, parsed_content: Dict) -> Document:
        """Apply Springer LNCS styles to parsed content"""
        self.create_new_document()
        self.figure_counter = 0
        self.table_counter = 0
        
        # Add title
        if parsed_content.get('title'):
            self.add_title(parsed_content['title'])
        
        # Add authors
        if parsed_content.get('authors'):
            self.add_authors(parsed_content['authors'])
        
        # Add abstract
        if parsed_content.get('abstract'):
            self.add_abstract(parsed_content['abstract'])
        
        # Add keywords
        if parsed_content.get('keywords'):
            self.add_keywords(parsed_content['keywords'])
        
        # Add sections
        if parsed_content.get('sections'):
            for section in parsed_content['sections']:
                self.add_section(section)
        
        # Add references
        if parsed_content.get('references'):
            self.add_references(parsed_content['references'])
        
        return self.document
    
    def add_title(self, title: str):
        """Add formatted title"""
        para = self.document.add_paragraph()
        para.style = self.document.styles['LNCS Title']
        run = para.add_run(title)
        run.font.name = self.styles.TITLE_FONT
        run.font.size = self.styles.TITLE_SIZE
        run.font.bold = True
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.changes_log.append("Title formatted (14pt, Bold, Center)")
    
    def add_authors(self, authors: List[Dict]):
        """Add formatted authors with affiliations"""
        if not authors:
            return
        
        # Author names
        names_para = self.document.add_paragraph()
        names_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for i, author in enumerate(authors):
            run = names_para.add_run(author.get('name', ''))
            run.font.name = self.styles.AUTHOR_FONT
            run.font.size = self.styles.AUTHOR_SIZE
            
            # Add superscript number
            sup_run = names_para.add_run(str(i + 1))
            sup_run.font.superscript = True
            sup_run.font.size = Pt(8)
            
            if i < len(authors) - 1:
                names_para.add_run(', ')
        
        self.changes_log.append("Authors formatted with superscript affiliations")
    
    def add_abstract(self, abstract_data):
        """Add formatted abstract"""
        abstract_text = ""
        if isinstance(abstract_data, dict):
            abstract_text = abstract_data.get('text', '')
        else:
            abstract_text = str(abstract_data)
        
        if not abstract_text:
            return
        
        para = self.document.add_paragraph()
        
        # Add "Abstract." prefix in bold
        prefix_run = para.add_run("Abstract. ")
        prefix_run.font.name = self.styles.ABSTRACT_FONT
        prefix_run.font.size = self.styles.ABSTRACT_SIZE
        prefix_run.font.bold = True
        
        # Add abstract text in italic
        text_run = para.add_run(abstract_text)
        text_run.font.name = self.styles.ABSTRACT_FONT
        text_run.font.size = self.styles.ABSTRACT_SIZE
        text_run.font.italic = True
        
        para.paragraph_format.space_after = Pt(12)
        
        self.changes_log.append('Abstract prefixed with "Abstract." (10pt, Italic)')
    
    def add_keywords(self, keywords: List[str]):
        """Add keywords"""
        if not keywords:
            return
        
        para = self.document.add_paragraph()
        
        # Add "Keywords:" prefix
        prefix_run = para.add_run("Keywords: ")
        prefix_run.font.name = self.styles.BODY_FONT
        prefix_run.font.size = self.styles.ABSTRACT_SIZE
        prefix_run.font.bold = True
        
        # Add keywords
        keyword_text = ' · '.join(keywords)
        text_run = para.add_run(keyword_text)
        text_run.font.name = self.styles.BODY_FONT
        text_run.font.size = self.styles.ABSTRACT_SIZE
        
        para.paragraph_format.space_after = Pt(18)
        
        self.changes_log.append(f"Keywords formatted ({len(keywords)} keywords)")
    
    def add_section(self, section: Dict):
        """Add a section with proper formatting"""
        section_type = section.get('type', '')
        title = section.get('title', '')
        content = section.get('content', '')
        
        # Add section heading
        heading_level = self._get_heading_level(section_type)
        self.add_heading(title, heading_level)
        
        # Add section content
        if content:
            for para_text in content.split('\n'):
                if para_text.strip():
                    self.add_body_paragraph(para_text.strip())
    
    def add_heading(self, text: str, level: int = 1):
        """Add a heading"""
        # Clean the text - remove existing numbers
        clean_text = re.sub(r'^\s*\d+\.?\s*', '', text)
        
        para = self.document.add_paragraph()
        run = para.add_run(clean_text)
        run.font.name = self.styles.BODY_FONT
        run.font.bold = True
        
        if level == 1:
            run.font.size = self.styles.HEADING1_SIZE
            para.paragraph_format.space_before = Pt(18)
        elif level == 2:
            run.font.size = self.styles.HEADING2_SIZE
            para.paragraph_format.space_before = Pt(12)
        else:
            run.font.size = self.styles.HEADING3_SIZE
            para.paragraph_format.space_before = Pt(9)
        
        para.paragraph_format.space_after = Pt(6)
    
    def add_body_paragraph(self, text: str):
        """Add a body paragraph"""
        para = self.document.add_paragraph()
        run = para.add_run(text)
        run.font.name = self.styles.BODY_FONT
        run.font.size = self.styles.BODY_SIZE
        
        para.paragraph_format.space_after = self.styles.PARAGRAPH_AFTER
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    
    def add_figure(self, image_path: str, caption: str = None, width_inches: float = 5.5):
        """Add a figure with caption"""
        self.figure_counter += 1
        
        # Add image
        if os.path.exists(image_path):
            img_para = self.document.add_paragraph()
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = img_para.add_run()
            run.add_picture(image_path, width=Inches(width_inches))
        
        # Add caption
        caption_text = caption if caption else f"Description of figure {self.figure_counter}"
        caption_para = self.document.add_paragraph()
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # "Fig. X." in bold
        fig_run = caption_para.add_run(f"Fig. {self.figure_counter}. ")
        fig_run.font.name = self.styles.CAPTION_FONT
        fig_run.font.size = self.styles.CAPTION_SIZE
        fig_run.font.bold = True
        
        # Caption text
        cap_run = caption_para.add_run(caption_text)
        cap_run.font.name = self.styles.CAPTION_FONT
        cap_run.font.size = self.styles.CAPTION_SIZE
        
        caption_para.paragraph_format.space_after = Pt(12)
        
        self.changes_log.append(f"Figure {self.figure_counter} inserted with caption")
        
        return self.figure_counter
    
    def add_references(self, references: List[str]):
        """Add references section"""
        if not references:
            return
        
        # Add References heading
        self.add_heading("References", 1)
        
        for i, ref in enumerate(references, 1):
            para = self.document.add_paragraph()
            
            # Reference number
            num_run = para.add_run(f"[{i}] ")
            num_run.font.name = self.styles.BODY_FONT
            num_run.font.size = Pt(9)
            
            # Reference text (clean any existing numbers)
            clean_ref = re.sub(r'^\s*\[?\d+\]?[\.\)]\s*', '', ref)
            ref_run = para.add_run(clean_ref)
            ref_run.font.name = self.styles.BODY_FONT
            ref_run.font.size = Pt(9)
            
            para.paragraph_format.space_after = Pt(3)
            para.paragraph_format.left_indent = Cm(0.5)
            para.paragraph_format.first_line_indent = Cm(-0.5)
        
        self.changes_log.append(f"References formatted: {len(references)} items")
    
    def insert_images_at_positions(self, images_data: List[Dict]):
        """Insert images at specified positions"""
        for img_info in images_data:
            self.add_figure(
                image_path=img_info.get('path', ''),
                caption=img_info.get('caption', ''),
                width_inches=img_info.get('width', 5.5)
            )
    
    def _get_heading_level(self, section_type: str) -> int:
        """Determine heading level from section type"""
        level_1_sections = ['introduction', 'related_work', 'methodology', 
                          'results', 'discussion', 'conclusion', 'references']
        
        if section_type in level_1_sections:
            return 1
        elif section_type == 'numbered_section':
            return 1
        return 2
    
    def save_document(self, output_path: str) -> bool:
        """Save the document"""
        try:
            self.document.save(output_path)
            return True
        except Exception as e:
            print(f"Error saving document: {e}")
            return False
    
    def get_changes_log(self) -> List[str]:
        """Get list of changes made"""
        return self.changes_log


def apply_springer_format(parsed_content: Dict, output_path: str) -> Dict:
    """Convenience function to apply Springer format and save"""
    applier = StyleApplier()
    applier.apply_springer_style(parsed_content)
    
    success = applier.save_document(output_path)
    
    return {
        'success': success,
        'output_path': output_path,
        'changes': applier.get_changes_log(),
        'figure_count': applier.figure_counter
    }
