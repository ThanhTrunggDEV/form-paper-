"""
Document Parser for Springer LNCS Formatter
Automatically detects and parses paper sections
"""

import re
from docx import Document
from typing import Dict, List, Optional, Tuple


class DocumentParser:
    """Parser for academic paper documents"""
    
    # Section detection patterns
    SECTION_PATTERNS = {
        'introduction': r'(?i)^\s*(1\.?\s*)?introduction\s*$',
        'related_work': r'(?i)^\s*(2\.?\s*)?(related\s*work|literature\s*review|background)\s*$',
        'methodology': r'(?i)^\s*(3\.?\s*)?(methodology|proposed\s*(method|approach|system)|method(s)?)\s*$',
        'results': r'(?i)^\s*(4\.?\s*)?(results?|experiments?|evaluation|experimental\s*results)\s*$',
        'discussion': r'(?i)^\s*(5\.?\s*)?discussion\s*$',
        'conclusion': r'(?i)^\s*(6\.?\s*)?(conclusion(s)?|summary|future\s*work)\s*$',
        'references': r'(?i)^\s*references?\s*$',
        'acknowledgments': r'(?i)^\s*(acknowledge?ments?)\s*$'
    }
    
    def __init__(self):
        self.document = None
        self.text_content = ""
        self.paragraphs = []
    
    def load_document(self, file_path: str) -> bool:
        """Load a Word document"""
        try:
            self.document = Document(file_path)
            self.paragraphs = [p.text.strip() for p in self.document.paragraphs if p.text.strip()]
            self.text_content = '\n'.join(self.paragraphs)
            return True
        except Exception as e:
            print(f"Error loading document: {e}")
            return False
    
    def load_from_text(self, text: str) -> bool:
        """Load from plain text"""
        try:
            self.text_content = text
            self.paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
            return True
        except Exception as e:
            print(f"Error loading text: {e}")
            return False
    
    def parse_sections(self) -> Dict:
        """Parse and detect all sections in the document"""
        result = {
            'title': self.extract_title(),
            'authors': self.extract_authors(),
            'abstract': self.extract_abstract(),
            'keywords': self.extract_keywords(),
            'sections': self.detect_sections(),
            'references': self.extract_references(),
            'raw_paragraphs': self.paragraphs
        }
        return result
    
    def extract_title(self) -> str:
        """Extract the title (usually first non-empty paragraph)"""
        if not self.paragraphs:
            return ""
        
        # Title is typically the first paragraph that's not too long
        for para in self.paragraphs[:3]:
            if len(para) < 200 and not para.lower().startswith('abstract'):
                return para
        return self.paragraphs[0] if self.paragraphs else ""
    
    def extract_authors(self) -> List[Dict]:
        """Extract author information"""
        authors = []
        
        # Look for author patterns in first few paragraphs
        author_pattern = r'([A-Z][a-z]+\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)'
        email_pattern = r'[\w\.-]+@[\w\.-]+\.\w+'
        
        for para in self.paragraphs[1:5]:
            # Skip if it looks like abstract
            if para.lower().startswith('abstract'):
                break
            
            # Find potential author names
            names = re.findall(author_pattern, para)
            emails = re.findall(email_pattern, para)
            
            for i, name in enumerate(names):
                author = {
                    'name': name,
                    'email': emails[i] if i < len(emails) else None,
                    'affiliation': None
                }
                authors.append(author)
        
        return authors
    
    def extract_abstract(self) -> Dict:
        """Extract abstract text"""
        abstract_text = ""
        abstract_started = False
        
        for i, para in enumerate(self.paragraphs):
            # Check if this is the abstract section
            if re.match(r'(?i)^\s*abstract\.?\s*', para):
                abstract_started = True
                # Remove "Abstract" prefix
                abstract_text = re.sub(r'(?i)^\s*abstract\.?\s*', '', para)
                continue
            
            if abstract_started:
                # Stop at keywords or next section
                if (re.match(r'(?i)^\s*keywords?\.?\s*:', para) or 
                    self._is_section_heading(para)):
                    break
                abstract_text += " " + para
        
        # Calculate word count
        word_count = len(abstract_text.split()) if abstract_text else 0
        
        return {
            'text': abstract_text.strip(),
            'word_count': word_count,
            'is_valid': 150 <= word_count <= 250
        }
    
    def extract_keywords(self) -> List[str]:
        """Extract keywords"""
        keywords = []
        
        for para in self.paragraphs:
            # Look for keywords pattern
            match = re.match(r'(?i)^\s*keywords?\.?\s*:?\s*(.+)$', para)
            if match:
                keyword_text = match.group(1)
                # Split by common delimiters
                keywords = [k.strip() for k in re.split(r'[,;·•]', keyword_text) if k.strip()]
                break
        
        return keywords
    
    def detect_sections(self) -> List[Dict]:
        """Detect all sections in the document"""
        sections = []
        current_section = None
        current_content = []
        
        for i, para in enumerate(self.paragraphs):
            section_type = self._get_section_type(para)
            
            if section_type:
                # Save previous section
                if current_section:
                    sections.append({
                        'type': current_section['type'],
                        'title': current_section['title'],
                        'content': '\n'.join(current_content),
                        'start_index': current_section['start_index']
                    })
                
                # Start new section
                current_section = {
                    'type': section_type,
                    'title': para,
                    'start_index': i
                }
                current_content = []
            elif current_section:
                current_content.append(para)
        
        # Add last section
        if current_section:
            sections.append({
                'type': current_section['type'],
                'title': current_section['title'],
                'content': '\n'.join(current_content),
                'start_index': current_section['start_index']
            })
        
        return sections
    
    def extract_references(self) -> List[str]:
        """Extract reference list"""
        references = []
        in_references = False
        
        for para in self.paragraphs:
            if re.match(self.SECTION_PATTERNS['references'], para):
                in_references = True
                continue
            
            if in_references:
                # Check for numbered reference pattern
                if re.match(r'^\s*\[?\d+\]?[\.\)]\s*', para) or para:
                    references.append(para)
        
        return references
    
    def _is_section_heading(self, text: str) -> bool:
        """Check if text is a section heading"""
        for pattern in self.SECTION_PATTERNS.values():
            if re.match(pattern, text):
                return True
        # Also check for numbered headings
        if re.match(r'^\s*\d+\.?\s+[A-Z]', text):
            return True
        return False
    
    def _get_section_type(self, text: str) -> Optional[str]:
        """Get section type from text"""
        for section_type, pattern in self.SECTION_PATTERNS.items():
            if re.match(pattern, text):
                return section_type
        
        # Check for generic numbered section
        if re.match(r'^\s*\d+\.?\s+[A-Z]', text):
            return 'numbered_section'
        
        return None
    
    def get_document_stats(self) -> Dict:
        """Get document statistics"""
        word_count = len(self.text_content.split())
        char_count = len(self.text_content)
        para_count = len(self.paragraphs)
        
        # Estimate page count (roughly 250-300 words per page)
        page_estimate = word_count / 275
        
        return {
            'word_count': word_count,
            'character_count': char_count,
            'paragraph_count': para_count,
            'estimated_pages': round(page_estimate, 1)
        }
    
    def find_image_insertion_points(self, keywords: List[str] = None) -> List[Dict]:
        """Find suitable points to insert images based on keywords"""
        if keywords is None:
            keywords = ['figure', 'fig.', 'image', 'diagram', 'chart', 'table', 
                       'shows', 'illustrates', 'demonstrates', 'presents']
        
        insertion_points = []
        
        for i, para in enumerate(self.paragraphs):
            for keyword in keywords:
                if keyword.lower() in para.lower():
                    insertion_points.append({
                        'paragraph_index': i,
                        'text': para[:100] + '...' if len(para) > 100 else para,
                        'keyword_found': keyword
                    })
                    break
        
        return insertion_points


def parse_docx_file(file_path: str) -> Dict:
    """Convenience function to parse a docx file"""
    parser = DocumentParser()
    if parser.load_document(file_path):
        return parser.parse_sections()
    return {}


def parse_text_content(text: str) -> Dict:
    """Convenience function to parse text content"""
    parser = DocumentParser()
    if parser.load_from_text(text):
        return parser.parse_sections()
    return {}
